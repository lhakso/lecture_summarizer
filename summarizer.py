import mlx_whisper
import re
import ollama
import subprocess
import time
from docx import Document
from docx.shared import Pt
from datetime import date
from pathlib import Path
import smtplib
from email.message import EmailMessage
from recorder import RecordingSession
import os
from tqdm import tqdm
from dotenv import load_dotenv

load_dotenv()
today = date.today()
target_word_count = "1200"
EMAIL_SENDER = os.getenv("EMAIL_SENDER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
EMAIL_RECEIVERS = os.getenv("EMAIL_RECEIVERS")


def entry():
    record_new = input("do you want to record a new lecture? y/n\n")

    if record_new != "y":
        print("using previous recording")
    else:
        session = RecordingSession(today=today)
        session.start_record()

    # start ollama in background
    ollama_process = subprocess.Popen(
        ["ollama", "serve"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )
    print("ollama server started!")
    time.sleep(2)  # give time to initialize

    try:
        summary = create_summary()
        path = create_doc(summary)
        send_email(path)

    finally:
        ollama_process.terminate()
        print("ollama server stopped.")


def split_into_sections(transcript: str) -> list:
    num_sections = 4
    words_per_section = len(transcript.split(" ")) // num_sections
    sentences = transcript.split(".")
    sections = []
    current_section = []
    word_count = 0

    for sentence in sentences:
        current_section.append(sentence + ".")
        word_count += len(sentence.split())

        if word_count >= words_per_section:
            sections.append(" ".join(current_section))
            current_section = []
            word_count = 0

    if current_section:
        sections.append(" ".join(current_section))

    return sections


def clean_filler_words(text: str) -> str:
    fillers = ["um", "uh", "you know", "like", "kind of", "sort of", "I mean"]
    pattern = r"\s*\b(?:" + "|".join(fillers) + r")[,.\s]*\b"
    cleaned_text = re.sub(pattern, " ", text, flags=re.IGNORECASE)
    return cleaned_text


def send_email(path: Path) -> None:
    email_sender = EMAIL_SENDER
    email_password = EMAIL_PASSWORD
    email_receivers = EMAIL_RECEIVERS
    file_path = Path(path)

    msg = EmailMessage()
    msg["Subject"] = "Lecture Summary"
    msg["From"] = email_sender
    msg["To"] = email_receivers
    msg.set_content(f"Notes for {today.strftime("%B %d, %Y")} lecture")

    with open(file_path, "rb") as file:
        msg.add_attachment(
            file.read(),
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=file_path.name,
        )

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(email_sender, email_password)
        smtp.send_message(msg)
    print("successfully sent email")


def transcribe_and_clean():
    file = Path(f"transcripts/transcript{today}.txt")

    if not file.exists():
        transcript = mlx_whisper.transcribe(
            f"audio/raw_audio{today}.wav", path_or_hf_repo="mlx-community/whisper-turbo"
        )["text"]

        with open(f"transcripts/transcript{today}.txt", "w") as file:
            file.write(transcript)
        print("successfully transcribed")
    else:
        print("skipped transcribe")
        with open(f"transcripts/transcript{today}.txt", "r") as file:
            transcript = file.read()

    cleaned_transcript = clean_filler_words(transcript)

    return cleaned_transcript


def create_summary() -> str:
    cleaned_transcript = transcribe_and_clean()
    split_transcript = split_into_sections(cleaned_transcript)
    print("cleaned and split transcript")

    summary = ""
    for section_index, section in enumerate(
        tqdm(split_transcript, desc="Summarizing Sections", unit="section")
    ):
        section_prompt = f"""
        Previous sections written so far:\n{summary}

        Now summarize the next part of the lecture in detail.
        - If there is a switch in topic from the previous section, include section header to indicate major topic.  
        - Do NOT introduce sections with unnecessary text like "Here is a summary."
        - Avoid phrases like "the professor said" and focus on the content.

        Here is the next section:

        {section}
"""

        response = ollama.chat(
            model="my-llama3-70b",
            # adjust message prompt as desired - second commented one is for testing
            messages=[{"role": "user", "content": section_prompt}],
            # messages=[{'role': 'user', 'content':f"please summarize this shortly: \n{cleaned_transcript}"}],
            stream=False,
        )
        summary += response["message"]["content"]
        print(f"section {section_index+1} done\n")
    return summary


def parse_output(text: str) -> list[tuple]:
    parsed = []
    for line in text.split("\n"):
        line = line.strip()

        if not line:
            parsed.append(("blank", ""))

        if line.startswith("**") and line.endswith("**"):
            heading = line.strip("*")
            parsed.append(("heading", heading))

        elif line.startswith("*"):
            bullet = line.lstrip("*")
            parsed.append(("bullet", bullet))

        else:
            parsed.append(("text", line))

    print("parsed output")

    return parsed


def create_doc(summary: str) -> str:

    parsed_summary = parse_output(summary)
    doc_path = os.getenv("DOC_PATH") + f"lecture_{today}.docx"

    doc = Document()

    doc.add_heading(f'Lecture Summary from {today.strftime("%B %d, %Y")}', level=1)
    for format_type, content in parsed_summary:
        if format_type == "blank":
            doc.add_paragraph("")

        elif format_type == "heading":
            doc.add_heading(content, level=3)

        elif format_type == "bullet":
            doc.add_paragraph(content, style="ListBullet")

        else:
            doc.add_paragraph(content)

    doc.save(doc_path)
    return doc_path


if __name__ == "__main__":
    entry()
